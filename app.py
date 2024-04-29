from flask import Flask, render_template, request, jsonify, redirect, url_for
from datetime import datetime
import requests
import base64
import io
import os
import PIL.Image
import textwrap
from docx import Document
from IPython.display import display, Markdown
import google.generativeai as genai
import openai
from bs4 import BeautifulSoup
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
import heapq
import requests
from gtts import gTTS
from youtube_transcript_api import YouTubeTranscriptApi

app = Flask(__name__)

api_key1 = "Your-Api-Key"
openai.api_key = api_key1

# Configure Generative AI API key
genai.configure(api_key="Your-Api-Key")

# Initialize Generative AI model
model = genai.GenerativeModel('gemini-pro-vision')



# Define the questions and their correct answers
questions_answers = [
    ("What is 2 + 2?", "4"),
    ("What is the output of 'print(3 * 5)'?", "15"),
    ("What is the result of '10 / 2'?", "5"),
    ("What will 'print(7 ** 2)' output?", "49"),
    ("What is the output of 'print('hello' + ' world')'?", "hello world"),
    ("What will 'print(len('python'))' output?", "6"),
    ("What is the output of 'print(10 % 3)'?", "1"),
    ("What is the value of 'x' if 'x = 10 // 3'?", "3"),
    ("What is the result of 'print(2 == 2)'?", "True"),
    ("What is the value of 'y' if 'x = 5; y = x * 2'?", "10"),
    ("What is the output of 'print(type(5))'?", "<class 'int'>"),
    ("What will 'print('hello'.upper())' output?", "HELLO"),
    ("What is the value of 'z' if 'x = 2; y = 3; z = x + y'?", "5"),
    ("What is the output of 'print('python'[2:])'?", "thon"),
    ("What will 'print('apple' * 3)' output?", "appleappleapple")
]


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/edu')
def edu():
    # Format date in day-month-year format
    for item in news_data:
        # Parse the ISO 8601 formatted date and convert it to day-month-year
        item['date'] = datetime.fromisoformat(item['date'].replace('Z', '+00:00')).strftime('%d-%m-%Y')

    return render_template('Edu.html', news=news_data)

# API endpoint for news data
url = 'https://oevortex-webscout.hf.space/api/news'
params = {
    'q': 'Technology',  # Query parameter
    'max_results': 10,   # Maximum number of results
    'safesearch': 'moderate',  # Safe search option
    'region': 'wt-wt'   # Region parameter
}

headers = {
    'accept': 'application/json'
}

# Get news data from API
response = requests.get(url, params=params, headers=headers)

if response.status_code == 200:
    news_data = response.json()['results']  # Access 'results' key
    # or do something with the data
else:
    print("Error:", response.status_code)

@app.route('/quiz')
def quizindex():
    return render_template('quiz-index.html', questions_answers=questions_answers)


@app.route('/submit_answers', methods=['POST'])
def submit_answers():
    user_answers = {key: request.form[key] for key in request.form}
    correct_answers_count = 0
    incorrect_answers_count = 0

    for question, correct_answer in questions_answers:
        user_answer = user_answers.get(question)
        if user_answer and user_answer.strip().lower() == correct_answer.lower():
            correct_answers_count += 1
        else:
            incorrect_answers_count += 1

    # Render the result template with the counts
    return render_template('quiz-result.html', correct_count=correct_answers_count, incorrect_count=incorrect_answers_count)


@app.route('/scan-solve')
def scanindex():
    return render_template('scan-solve.html')

@app.route('/upload', methods=['POST'])
def upload():
    try:
        # Get the image data from the request
        data = request.json
        image_data = data.get('image')

        # Decode the base64 image data
        image_data = image_data.split(',')[1]
        image_binary = base64.b64decode(image_data)

        # Save the image to a file on your local PC
        img_path = 'question.png'
        with open(img_path, 'wb') as img_file:
            img_file.write(image_binary)

        # Load the image
        img = PIL.Image.open(io.BytesIO(image_binary))

        # Use Generative AI model to generate text from the image
        response = model.generate_content(["Give me a simple answer for this question", img], stream=True)
        response.resolve()

        # Store response text in a Notepad file
        notepad_file_path = 'solution.txt'
        with open(notepad_file_path, 'w') as file:
            file.write(response.text)

        return jsonify({'message': 'Response text stored in Notepad file.'})

    except Exception as e:
        return jsonify({'error': str(e)}), 500



@app.route('/transcribe')
def speechindex():
    return render_template('transcript.html')

@app.route('/store_transcription', methods=['POST'])
def store_transcription():
    data = request.json
    transcription = data['transcription']

    # Store transcription in a Word file
    document = Document()
    document.add_paragraph(transcription)
    document.save('transcription.docx')

    return 'Transcription stored successfully'



@app.route('/speech-summary')
def get_summary():
    # Read the text from the .docx file
    docx_file = 'transcription.docx'
    doc_text = read_docx(docx_file)

    # Generate summary
    summary = generate_summary(doc_text)

    return summary


def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)


def to_markdown(text):
    text = text.replace('\u2022', '  *')
    return Markdown(textwrap.indent(text, '> ', predicate=lambda _: True))


def generate_summary(note_text):
    model = genai.GenerativeModel('gemini-pro')
    rply = model.generate_content("summarize my the given text and dont add any new words just summarize the given text if needed add one or two lines on your own and it should be short " + note_text)
    return render_template('transcript-result.html', script_Summary=rply.text)


@app.route('/road')
def road():
    return render_template('road.html')

@app.route('/road-result', methods=['POST'])
def generate_road_plan():
    subject = request.form['subject']
    
    # Define the prompt
    prompt = f"Generate roadmap for {subject} ."

    # Generate the completion
    completion = openai.Completion.create(
        model="gpt-3.5-turbo-instruct",  # Use the model that supports stream=True
        prompt=prompt,
        max_tokens=150,
    )

    # Extract and return the response
    roadmap_plan = completion.choices[0].text.strip()
    return render_template('road-result.html', result=roadmap_plan)


@app.route('/codemaster')
def code():
    return render_template('code.html')

@app.route('/generate', methods=['POST'])
def generate_html():
    prompt = request.form['prompt']
    model = genai.GenerativeModel('gemini-pro')
    rply = model.generate_content("Generate both html and css code in single file for " + prompt + " with colorful CSS background and more attractive CSS, I need only code and no explanation")
    html_content = rply.text

    with open("templates/index1.html", "w") as file:
        file.write(html_content)

    return redirect(url_for('output'))


@app.route('/output')
def output():
    return render_template('index1.html')

@app.route('/chatbot')
def chatbot():
    return render_template('chatbot.html')

@app.route('/ask', methods=['POST'])
def ask_question():
    user_question = request.form['question']

    # Generate response from OpenAI's GPT-3.5 model
    prompt = f"Answer for my question in short: {user_question}"
    completion = openai.Completion.create(
        model="gpt-3.5-turbo-instruct",
        prompt=prompt,
        max_tokens=150,
        stop=None,
        temperature=0.7
    )
    chatgpt_response = completion.choices[0].text.strip()

    return chatgpt_response

def generate_notesummary(note_text):
    model = genai.GenerativeModel('gemini-pro')
    rply = model.generate_content("summarize my notes"+note_text)
    to_markdown(rply.text)
    return rply.text
# Route for home page with file upload form
@app.route('/note', methods=['GET', 'POST'])
def note():
    if request.method == 'POST':
        # Check if a file was uploaded
        if 'Note_file' not in request.files:
            return render_template('error.html', message='No file part')

        file = request.files['Note_file']

        # Check if the file is empty
        if file.filename == '':
            return render_template('error.html', message='No selected file')

        # Check if the file is of allowed type
        if file and file.filename.endswith('.txt'):
            # Read the file content
            note_text = file.read().decode('utf-8')

            # Generate summary
            summary_text = generate_notesummary(note_text)
            
            # Render the result template with summary
            return render_template('note-result.html', summary_text=summary_text)

        else:
            return render_template('error.html', message='Invalid file type. Please upload a text file')

    return render_template('note.html')
	

def scrape_website(url):
    try:
        # Send a GET request to the URL
        response = requests.get(url)
        response.raise_for_status()  # Raise an exception for HTTP errors

        # Parse the HTML content of the webpage
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract all text from the webpage
        text = soup.get_text()

        return text
    except requests.exceptions.RequestException as e:
        print("Error fetching the webpage:", e)
        return None
    

def summarize_text(text, num_sentences=3):
    sentences = sent_tokenize(text)
    word_frequencies = FreqDist()
    for word in nltk.word_tokenize(text):
        if word.lower() not in stopwords.words('english'):
            word_frequencies[word.lower()] += 1

    most_frequent_words = max(word_frequencies.values())
    for word in word_frequencies.keys():
        word_frequencies[word] = (word_frequencies[word] / most_frequent_words)

    sentence_scores = {}
    for sentence in sentences:
        for word in nltk.word_tokenize(sentence.lower()):
            if word in word_frequencies.keys():
                if len(sentence.split(' ')) < 30:
                    if sentence not in sentence_scores.keys():
                        sentence_scores[sentence] = word_frequencies[word]
                    else:
                        sentence_scores[sentence] += word_frequencies[word]

    summary_sentences = heapq.nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    summary = ' '.join(summary_sentences)

    return summary

@app.route('/web')
def web():
    return render_template('web-summary.html')

@app.route('/web_summary', methods=['POST'])
def web_summary():
    url = request.form['url']
    scraped_text = scrape_website(url)

    if scraped_text:
        websummary = summarize_text(scraped_text)
        print(websummary)
        tts = gTTS(text=websummary, lang='en')
        output_file_audio = "static/outputsummary.mp3"
        tts.save(output_file_audio)
        return render_template('web-result.html', websummary=websummary, audio_file=output_file_audio)
    else:
        return "No text scraped from the website."

@app.route('/youtube-summary')
def yt():
    return render_template('youtube.html')
	


	
	
@app.route('/summarize', methods=['POST'])
def summarize():
    video_link = request.form['videoLink']
    video_id = video_link.split('=')[1]

    transcript = YouTubeTranscriptApi.get_transcript(video_id)
    result = ""	
    for i in transcript:
        result += ' ' + i['text']

    text_summary = summarize_text(result)
    tts = gTTS(text=text_summary, lang='en')
    tts.save('ytoutput.mp3')

    return render_template('youtube-result.html', summary=text_summary)




@app.route('/play_audio')
def play_audio():
    return send_file('ytoutput.mp3', mimetype='audio/mpeg')




@app.route('/ai-writer')
def writer():
    return render_template('letter.html')
	
@app.route('/generate-letter', methods=['POST'])
def generate_letter():
    prompt = request.form['prompt']
    model = genai.GenerativeModel('gemini-pro')
    rply = model.generate_content(" Generate an " + prompt + " with proper alignment ")
    letter = rply.text

    with open("letter.txt", "w") as file:
        file.write(letter)
    return render_template('index.html')



if __name__ == '__main__':
    app.run(debug=True)